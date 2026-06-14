"""Tests for the minimal permit web interface."""

from datetime import timedelta
from pathlib import Path
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group
from django.core.files.base import ContentFile
from django.core.management import call_command
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils.http import urlencode
from django.utils import timezone

from approvals.models import ApprovalAction
from audit.models import AuditLog
from documents.models import DocumentTemplate, GeneratedDocument
from permits.models import (
    Equipment,
    Hazard,
    Permit,
    PermitParticipant,
    PermitParticipantRole,
    PermitStatus,
    Personnel,
    PersonnelGroup,
    SafetyMeasure,
    WorkArea,
    WorkType,
)
from users.roles import ROLE_CHIEF, ROLE_MASTER, ROLE_OPERATOR


class AuthenticationViewTests(TestCase):
    """Checks that users can log in through the web UI."""

    def test_user_can_log_in_with_valid_credentials(self):
        user_model = get_user_model()
        user_model.objects.create_user(username="operator-login", password="secret-pass")

        response = self.client.post(
            reverse("login"),
            data={"username": "operator-login", "password": "secret-pass"},
        )

        self.assertRedirects(response, reverse("permits:list"))

    def test_login_page_renders(self):
        response = self.client.get(reverse("login"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Login")


class PermitViewTests(TestCase):
    """Checks page availability and basic permit form behaviour."""

    @classmethod
    def setUpTestData(cls):
        call_command("setup_roles", verbosity=0)
        user_model = get_user_model()
        cls.operator = user_model.objects.create_user(username="operator", password="pass")
        cls.master = user_model.objects.create_user(username="master", password="pass")
        cls.chief = user_model.objects.create_user(username="chief", password="pass")
        cls.plain_user = user_model.objects.create_user(username="plain", password="pass")
        cls.manager = user_model.objects.create_user(username="manager", password="pass")
        cls.supervisor = user_model.objects.create_user(username="supervisor", password="pass")
        cls.operator.groups.add(Group.objects.get(name=ROLE_OPERATOR))
        cls.master.groups.add(Group.objects.get(name=ROLE_MASTER))
        cls.chief.groups.add(Group.objects.get(name=ROLE_CHIEF))
        cls.work_area = WorkArea.objects.create(name="Workshop 1")
        cls.equipment = Equipment.objects.create(
            name="Pump A",
            code="P-100",
            work_area=cls.work_area,
        )
        cls.work_type = WorkType.objects.create(name="Inspection")
        cls.hazard = Hazard.objects.create(name="Pressure")
        cls.safety_measure = SafetyMeasure.objects.create(name="Lockout")
        cls.personnel_group = PersonnelGroup.objects.create(name="Мастера")
        cls.personnel_manager = Personnel.objects.create(
            full_name="Иванов Иван Иванович",
            personnel_number="P-100",
            position="мастер",
            group=cls.personnel_group,
            work_area=cls.work_area,
        )
        cls.personnel_performer = Personnel.objects.create(
            full_name="Петров Пётр Петрович",
            personnel_number="P-200",
            position="слесарь",
            group=cls.personnel_group,
            work_area=cls.work_area,
        )

    def setUp(self):
        self.client.force_login(self.operator)

    def make_permit(self, number="PT-WEB-001", status=PermitStatus.DRAFT):
        starts_at = timezone.now() + timedelta(days=1)
        return Permit.objects.create(
            number=number,
            status=status,
            work_starts_at=starts_at,
            work_ends_at=starts_at + timedelta(hours=8),
            work_location="Workshop 1",
            responsible_manager_text="Manual manager",
            work_producer_text="Manual producer",
            work_nature_text="Manual work nature",
            additional_conditions="Manual additional conditions",
            additional_safety_notes="Manual additional safety notes",
            work_area=self.work_area,
            equipment=self.equipment,
            work_type=self.work_type,
            work_description="Inspect and repair valve",
            responsible_manager=self.manager,
            work_supervisor=self.supervisor,
            created_by=self.operator,
        )

    def permit_form_data(self, number="PT-WEB-NEW"):
        starts_at = timezone.now() + timedelta(days=2)
        ends_at = starts_at + timedelta(hours=4)
        return {
            "number": number,
            "work_starts_at": starts_at.strftime("%Y-%m-%dT%H:%M"),
            "work_ends_at": ends_at.strftime("%Y-%m-%dT%H:%M"),
            "work_location": "Pump station",
            "responsible_manager_text": "Manual form manager",
            "work_producer_text": "Manual form producer",
            "work_nature_text": "Manual form work nature",
            "additional_conditions": "Manual form additional conditions",
            "additional_safety_notes": "Manual form additional safety notes",
            "work_area": self.work_area.pk,
            "equipment": self.equipment.pk,
            "work_type": self.work_type.pk,
            "hazards": [self.hazard.pk],
            "safety_measures": [self.safety_measure.pk],
            "work_description": "Replace gasket",
            "responsible_manager": self.manager.pk,
            "work_supervisor": self.supervisor.pk,
            "participants-TOTAL_FORMS": "2",
            "participants-INITIAL_FORMS": "0",
            "participants-MIN_NUM_FORMS": "0",
            "participants-MAX_NUM_FORMS": "1000",
            "participants-0-role": PermitParticipantRole.RESPONSIBLE_MANAGER,
            "participants-0-personnel": self.personnel_manager.pk,
            "participants-0-manual_name": "",
            "participants-0-note": "from directory",
            "participants-0-sort_order": "1",
            "participants-1-role": PermitParticipantRole.PERFORMER,
            "participants-1-personnel": "",
            "participants-1-manual_name": "Manual performer",
            "participants-1-note": "manual participant",
            "participants-1-sort_order": "2",
        }

    def test_permit_list_page_shows_table_and_permit(self):
        permit = self.make_permit()

        response = self.client.get(reverse("permits:list"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, permit.number)
        self.assertContains(response, permit.work_location)
        self.assertContains(response, "status-badge status-draft")
        self.assertContains(response, reverse("permits:detail", kwargs={"pk": permit.pk}))

    def test_dashboard_page_available_for_authenticated_user(self):
        permit = self.make_permit(number="PT-DASHBOARD-RECENT", status=PermitStatus.DRAFT)

        response = self.client.get(reverse("permits:dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Панель управления")
        self.assertContains(response, "Количество нарядов по статусам")
        self.assertContains(response, "Последние наряды")
        self.assertContains(response, permit.number)
        self.assertContains(response, "status-badge status-draft")
        self.assertContains(response, "Создать наряд")
        self.assertContains(response, "Список нарядов")

    def test_dashboard_redirects_anonymous_user_to_login(self):
        self.client.logout()

        response = self.client.get(reverse("permits:dashboard"))

        self.assertRedirects(
            response,
            f"{reverse('login')}?next={reverse('permits:dashboard')}",
        )


    def test_personnel_search_finds_active_personnel(self):
        url = reverse("personnel_search") + "?" + urlencode({"q": "Иванов"})

        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["results"][0]["id"], self.personnel_manager.pk)
        self.assertIn("Иванов Иван Иванович", response.json()["results"][0]["label"])

    def test_dashboard_operator_sees_returned_permits_waiting_for_action(self):
        returned = self.make_permit(number="PT-DASHBOARD-RETURNED", status=PermitStatus.RETURNED)
        submitted = self.make_permit(number="PT-DASHBOARD-SUBMITTED", status=PermitStatus.SUBMITTED)

        response = self.client.get(reverse("permits:dashboard"))

        self.assertEqual(response.status_code, 200)
        pending_permits = list(response.context["pending_permits"])
        self.assertIn(returned, pending_permits)
        self.assertNotIn(submitted, pending_permits)

    def test_dashboard_master_sees_submitted_permits_waiting_for_action(self):
        self.client.force_login(self.master)
        submitted = self.make_permit(number="PT-DASHBOARD-MASTER", status=PermitStatus.SUBMITTED)
        returned = self.make_permit(number="PT-DASHBOARD-OPERATOR", status=PermitStatus.RETURNED)

        response = self.client.get(reverse("permits:dashboard"))

        self.assertEqual(response.status_code, 200)
        pending_permits = list(response.context["pending_permits"])
        self.assertIn(submitted, pending_permits)
        self.assertNotIn(returned, pending_permits)

    def test_dashboard_chief_sees_master_approved_permits_waiting_for_action(self):
        self.client.force_login(self.chief)
        master_approved = self.make_permit(
            number="PT-DASHBOARD-CHIEF",
            status=PermitStatus.APPROVED_BY_MASTER,
        )
        submitted = self.make_permit(number="PT-DASHBOARD-MASTER-ONLY", status=PermitStatus.SUBMITTED)

        response = self.client.get(reverse("permits:dashboard"))

        self.assertEqual(response.status_code, 200)
        pending_permits = list(response.context["pending_permits"])
        self.assertIn(master_approved, pending_permits)
        self.assertNotIn(submitted, pending_permits)

    def test_permit_detail_page_shows_core_sections(self):
        permit = self.make_permit()
        PermitParticipant.objects.create(
            permit=permit,
            role=PermitParticipantRole.RESPONSIBLE_MANAGER,
            personnel=self.personnel_manager,
            note="detail",
            sort_order=1,
        )
        PermitParticipant.objects.create(
            permit=permit,
            role=PermitParticipantRole.PERFORMER,
            manual_name="Manual detail performer",
            sort_order=2,
        )
        action = ApprovalAction.objects.create(
            permit=permit,
            actor=self.operator,
            action="submit",
            old_status=PermitStatus.DRAFT,
            new_status=PermitStatus.SUBMITTED,
            comment="submitted",
        )
        template = DocumentTemplate.objects.create(
            name="Template",
            document_type="permit-web",
            version="1",
            file="templates/template.docx",
            uploaded_by=self.operator,
        )
        document = GeneratedDocument.objects.create(
            permit=permit,
            template=template,
            file_docx="generated/permit.docx",
            generated_by=self.operator,
        )

        AuditLog.objects.create(
            user=self.operator,
            action="permit.updated",
            object_type="Permit",
            object_id=str(permit.pk),
            details={
                "changes": {
                    "work_location": {
                        "old": "Old location",
                        "new": permit.work_location,
                    }
                }
            },
        )

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Основные поля")
        self.assertContains(response, permit.number)
        self.assertContains(response, "status-badge status-draft")
        self.assertContains(response, action.comment)
        self.assertContains(response, document.file_docx.name)
        self.assertContains(response, "История изменений наряда")
        self.assertContains(response, "Изменено")
        self.assertContains(response, "work_location")
        self.assertContains(response, "было=Old location")
        self.assertContains(response, f"стало={permit.work_location}")
        self.assertContains(response, "Отправить на проверку")
        self.assertContains(response, "Действия")
        self.assertContains(response, "Комментарий для Отправить на проверку")
        self.assertContains(response, "Участники и ответственные")
        self.assertContains(response, "Ответственные руководители")
        self.assertContains(response, "Иванов Иван Иванович")
        self.assertContains(response, "Manual detail performer")

    def test_create_permit_page_creates_draft_permit(self):
        response = self.client.post(reverse("permits:create"), data=self.permit_form_data())

        permit = Permit.objects.get(number="PT-WEB-NEW")
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.status, PermitStatus.DRAFT)
        self.assertEqual(permit.created_by, self.operator)
        self.assertEqual(permit.work_area, self.work_area)
        self.assertEqual(permit.equipment, self.equipment)
        self.assertEqual(permit.work_type, self.work_type)
        self.assertIn(self.hazard, permit.hazards.all())
        self.assertIn(self.safety_measure, permit.safety_measures.all())
        self.assertEqual(permit.responsible_manager_text, "Manual form manager")
        self.assertEqual(permit.work_producer_text, "Manual form producer")
        self.assertEqual(permit.work_nature_text, "Manual form work nature")
        self.assertEqual(permit.additional_conditions, "Manual form additional conditions")
        self.assertEqual(permit.additional_safety_notes, "Manual form additional safety notes")
        self.assertEqual(permit.participants.count(), 2)
        self.assertTrue(
            permit.participants.filter(
                role=PermitParticipantRole.RESPONSIBLE_MANAGER,
                personnel=self.personnel_manager,
            ).exists()
        )
        self.assertTrue(
            permit.participants.filter(
                role=PermitParticipantRole.PERFORMER,
                manual_name="Manual performer",
            ).exists()
        )

    def test_create_permit_page_writes_audit_log(self):
        response = self.client.post(reverse("permits:create"), data=self.permit_form_data("PT-WEB-AUDIT-NEW"))

        permit = Permit.objects.get(number="PT-WEB-AUDIT-NEW")
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        audit_log = AuditLog.objects.get(action="permit.created", object_id=str(permit.pk))
        self.assertEqual(audit_log.user, self.operator)
        self.assertEqual(audit_log.object_type, "Permit")
        self.assertEqual(audit_log.details["old_values"], {})
        self.assertEqual(audit_log.details["new_values"]["number"], "PT-WEB-AUDIT-NEW")
        self.assertEqual(audit_log.details["new_values"]["work_area"]["label"], "Workshop 1")
        self.assertEqual(audit_log.details["new_values"]["hazards"], [{"id": self.hazard.pk, "label": "Pressure"}])
        self.assertIsNotNone(audit_log.created_at)

    def test_edit_permit_page_updates_draft_permit(self):
        permit = self.make_permit()
        data = self.permit_form_data(number=permit.number)
        data["work_location"] = "Updated location"
        data["work_nature_text"] = "Updated manual work nature"

        response = self.client.post(reverse("permits:edit", kwargs={"pk": permit.pk}), data=data)

        permit.refresh_from_db()
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.work_location, "Updated location")
        self.assertEqual(permit.work_nature_text, "Updated manual work nature")
        self.assertEqual(permit.participants.count(), 2)

    def test_edit_permit_page_updates_and_deletes_participants(self):
        permit = self.make_permit()
        responsible = PermitParticipant.objects.create(
            permit=permit,
            role=PermitParticipantRole.RESPONSIBLE_MANAGER,
            personnel=self.personnel_manager,
            note="old note",
            sort_order=1,
        )
        performer = PermitParticipant.objects.create(
            permit=permit,
            role=PermitParticipantRole.PERFORMER,
            manual_name="Old performer",
            sort_order=2,
        )
        data = self.permit_form_data(number=permit.number)
        data.update(
            {
                "participants-TOTAL_FORMS": "2",
                "participants-INITIAL_FORMS": "2",
                "participants-0-id": responsible.pk,
                "participants-0-role": PermitParticipantRole.RESPONSIBLE_MANAGER,
                "participants-0-personnel": self.personnel_manager.pk,
                "participants-0-manual_name": "",
                "participants-0-note": "updated note",
                "participants-0-sort_order": "1",
                "participants-1-id": performer.pk,
                "participants-1-role": PermitParticipantRole.PERFORMER,
                "participants-1-personnel": "",
                "participants-1-manual_name": "Old performer",
                "participants-1-note": "",
                "participants-1-sort_order": "2",
                "participants-1-DELETE": "on",
            }
        )

        response = self.client.post(reverse("permits:edit", kwargs={"pk": permit.pk}), data=data)

        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        responsible.refresh_from_db()
        self.assertEqual(responsible.note, "updated note")
        self.assertFalse(PermitParticipant.objects.filter(pk=performer.pk).exists())

    def test_invalid_permit_form_does_not_save_participants(self):
        data = self.permit_form_data(number="")

        response = self.client.post(reverse("permits:create"), data=data)

        self.assertEqual(response.status_code, 200)
        self.assertFalse(Permit.objects.filter(work_location="Pump station").exists())
        self.assertFalse(PermitParticipant.objects.filter(manual_name="Manual performer").exists())

    def test_empty_participant_formset_row_is_ignored(self):
        data = self.permit_form_data(number="PT-WEB-EMPTY-PARTICIPANT")
        data.update(
            {
                "participants-TOTAL_FORMS": "1",
                "participants-INITIAL_FORMS": "0",
                "participants-0-role": PermitParticipantRole.PERFORMER,
                "participants-0-personnel": "",
                "participants-0-manual_name": "",
                "participants-0-note": "",
                "participants-0-sort_order": "0",
            }
        )

        response = self.client.post(reverse("permits:create"), data=data)

        permit = Permit.objects.get(number="PT-WEB-EMPTY-PARTICIPANT")
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.participants.count(), 0)

    def test_edit_permit_page_writes_changed_fields_audit_log(self):
        permit = self.make_permit(number="PT-WEB-AUDIT-EDIT")
        data = self.permit_form_data(number=permit.number)
        data["work_location"] = "Audited updated location"

        response = self.client.post(reverse("permits:edit", kwargs={"pk": permit.pk}), data=data)

        permit.refresh_from_db()
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        audit_log = AuditLog.objects.get(action="permit.updated", object_id=str(permit.pk))
        self.assertEqual(audit_log.user, self.operator)
        self.assertEqual(audit_log.object_type, "Permit")
        self.assertIn("work_location", audit_log.details["changes"])
        self.assertEqual(audit_log.details["old_values"]["work_location"], "Workshop 1")
        self.assertEqual(audit_log.details["new_values"]["work_location"], "Audited updated location")
        self.assertEqual(audit_log.details["changes"]["work_location"]["old"], "Workshop 1")
        self.assertEqual(
            audit_log.details["changes"]["work_location"]["new"],
            "Audited updated location",
        )
        self.assertIsNotNone(audit_log.created_at)

    def test_edit_permit_page_does_not_write_audit_log_without_actual_changes(self):
        starts_at = timezone.now().replace(second=0, microsecond=0) + timedelta(days=1)
        permit = Permit.objects.create(
            number="PT-WEB-AUDIT-NOCHANGE",
            status=PermitStatus.DRAFT,
            work_starts_at=starts_at,
            work_ends_at=starts_at + timedelta(hours=8),
            work_location="Workshop 1",
            responsible_manager_text="Manual manager",
            work_producer_text="Manual producer",
            work_nature_text="Manual work nature",
            additional_conditions="Manual additional conditions",
            additional_safety_notes="Manual additional safety notes",
            work_area=self.work_area,
            equipment=self.equipment,
            work_type=self.work_type,
            work_description="Inspect and repair valve",
            responsible_manager=self.manager,
            work_supervisor=self.supervisor,
            created_by=self.operator,
        )
        data = {
            "number": permit.number,
            "work_starts_at": timezone.localtime(permit.work_starts_at).strftime("%Y-%m-%dT%H:%M"),
            "work_ends_at": timezone.localtime(permit.work_ends_at).strftime("%Y-%m-%dT%H:%M"),
            "work_location": permit.work_location,
            "responsible_manager_text": permit.responsible_manager_text,
            "work_producer_text": permit.work_producer_text,
            "work_nature_text": permit.work_nature_text,
            "additional_conditions": permit.additional_conditions,
            "additional_safety_notes": permit.additional_safety_notes,
            "work_area": self.work_area.pk,
            "equipment": self.equipment.pk,
            "work_type": self.work_type.pk,
            "hazards": [],
            "safety_measures": [],
            "work_description": permit.work_description,
            "responsible_manager": self.manager.pk,
            "work_supervisor": self.supervisor.pk,
            "participants-TOTAL_FORMS": "0",
            "participants-INITIAL_FORMS": "0",
            "participants-MIN_NUM_FORMS": "0",
            "participants-MAX_NUM_FORMS": "1000",
        }

        response = self.client.post(reverse("permits:edit", kwargs={"pk": permit.pk}), data=data)

        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertFalse(AuditLog.objects.filter(action="permit.updated", object_id=str(permit.pk)).exists())

    def test_edit_permit_page_rejects_approved_permit(self):
        permit = self.make_permit(status=PermitStatus.APPROVED_BY_MASTER)

        response = self.client.get(reverse("permits:edit", kwargs={"pk": permit.pk}))

        self.assertEqual(response.status_code, 403)

    def test_operator_can_submit_draft_permit_with_comment(self):
        permit = self.make_permit(number="PT-WEB-ACTION-001", status=PermitStatus.DRAFT)

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "submit"}),
            data={"comment": "ready for review"},
        )

        permit.refresh_from_db()
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.status, PermitStatus.SUBMITTED)
        action = ApprovalAction.objects.get(permit=permit, action="submit")
        self.assertEqual(action.actor, self.operator)
        self.assertEqual(action.comment, "ready for review")
        self.assertTrue(
            AuditLog.objects.filter(
                object_id=str(permit.pk),
                action="permit.submit",
                details__comment="ready for review",
            ).exists()
        )

    def test_master_can_approve_submitted_permit_with_comment(self):
        permit = self.make_permit(number="PT-WEB-ACTION-002", status=PermitStatus.SUBMITTED)
        self.client.force_login(self.master)

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "approve_by_master"}),
            data={"comment": "master approval"},
        )

        permit.refresh_from_db()
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.status, PermitStatus.APPROVED_BY_MASTER)
        action = ApprovalAction.objects.get(permit=permit, action="approve_by_master")
        self.assertEqual(action.actor, self.master)
        self.assertEqual(action.comment, "master approval")
        self.assertTrue(
            AuditLog.objects.filter(
                object_id=str(permit.pk),
                action="permit.approve_by_master",
            ).exists()
        )

    def test_master_can_return_submitted_permit_with_comment(self):
        permit = self.make_permit(number="PT-WEB-ACTION-RETURN", status=PermitStatus.SUBMITTED)
        self.client.force_login(self.master)

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "return"}),
            data={"comment": "needs rework"},
        )

        permit.refresh_from_db()
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.status, PermitStatus.RETURNED)
        action = ApprovalAction.objects.get(permit=permit, action="return")
        self.assertEqual(action.actor, self.master)
        self.assertEqual(action.comment, "needs rework")
        self.assertTrue(
            AuditLog.objects.filter(
                object_id=str(permit.pk),
                action="permit.return",
                details__comment="needs rework",
            ).exists()
        )

    def test_chief_can_approve_master_approved_permit_with_comment(self):
        permit = self.make_permit(
            number="PT-WEB-ACTION-003",
            status=PermitStatus.APPROVED_BY_MASTER,
        )
        self.client.force_login(self.chief)

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "approve_by_chief"}),
            data={"comment": "chief approval"},
        )

        permit.refresh_from_db()
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertEqual(permit.status, PermitStatus.APPROVED_BY_CHIEF)
        action = ApprovalAction.objects.get(permit=permit, action="approve_by_chief")
        self.assertEqual(action.actor, self.chief)
        self.assertEqual(action.comment, "chief approval")
        self.assertTrue(
            AuditLog.objects.filter(
                object_id=str(permit.pk),
                action="permit.approve_by_chief",
            ).exists()
        )

    def test_disallowed_web_action_is_forbidden(self):
        permit = self.make_permit(number="PT-WEB-ACTION-004", status=PermitStatus.DRAFT)

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "approve_by_master"}),
            data={"comment": "not allowed"},
        )

        permit.refresh_from_db()
        self.assertEqual(response.status_code, 403)
        self.assertEqual(permit.status, PermitStatus.DRAFT)
        self.assertFalse(ApprovalAction.objects.filter(permit=permit).exists())
        self.assertFalse(AuditLog.objects.filter(object_id=str(permit.pk)).exists())

    def test_user_without_role_cannot_perform_action(self):
        permit = self.make_permit(number="PT-WEB-ACTION-NOROLE", status=PermitStatus.DRAFT)
        self.client.force_login(self.plain_user)

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "submit"}),
            data={"comment": "I should not be allowed"},
        )

        permit.refresh_from_db()
        self.assertEqual(response.status_code, 403)
        self.assertEqual(permit.status, PermitStatus.DRAFT)
        self.assertFalse(ApprovalAction.objects.filter(permit=permit).exists())
        self.assertFalse(AuditLog.objects.filter(object_id=str(permit.pk)).exists())

    def test_unauthenticated_user_cannot_perform_action(self):
        permit = self.make_permit(number="PT-WEB-ACTION-ANON", status=PermitStatus.DRAFT)
        self.client.logout()

        response = self.client.post(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "submit"}),
            data={"comment": "anonymous"},
        )

        permit.refresh_from_db()
        self.assertEqual(response.status_code, 302)
        self.assertIn("/accounts/login/", response["Location"])
        self.assertEqual(permit.status, PermitStatus.DRAFT)
        self.assertFalse(ApprovalAction.objects.filter(permit=permit).exists())
        self.assertFalse(AuditLog.objects.filter(object_id=str(permit.pk)).exists())

    def test_action_cannot_be_performed_with_get(self):
        permit = self.make_permit(number="PT-WEB-ACTION-GET", status=PermitStatus.DRAFT)

        response = self.client.get(
            reverse("permits:action", kwargs={"pk": permit.pk, "action": "submit"}),
        )

        permit.refresh_from_db()
        self.assertEqual(response.status_code, 405)
        self.assertEqual(permit.status, PermitStatus.DRAFT)
        self.assertFalse(ApprovalAction.objects.filter(permit=permit).exists())
        self.assertFalse(AuditLog.objects.filter(object_id=str(permit.pk)).exists())

    def test_detail_hides_disallowed_action_buttons(self):
        permit = self.make_permit(number="PT-WEB-ACTION-005", status=PermitStatus.DRAFT)

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Отправить на проверку")
        self.assertNotContains(response, "Согласовать мастером")
        self.assertNotContains(response, "Утвердить начальником")


    def test_master_sees_submitted_status_actions(self):
        permit = self.make_permit(number="PT-WEB-ACTION-006", status=PermitStatus.SUBMITTED)
        self.client.force_login(self.master)

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Вернуть на доработку")
        self.assertContains(response, "Согласовать мастером")
        self.assertContains(response, "Отклонить")
        self.assertNotContains(response, "Отправить на проверку")

    def test_chief_sees_master_approved_status_actions(self):
        permit = self.make_permit(
            number="PT-WEB-ACTION-007",
            status=PermitStatus.APPROVED_BY_MASTER,
        )
        self.client.force_login(self.chief)

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Вернуть на доработку")
        self.assertContains(response, "Утвердить начальником")
        self.assertContains(response, "Отклонить")
        self.assertNotContains(response, "Согласовать мастером")

    def test_operator_sees_close_action_for_chief_approved_permit(self):
        permit = self.make_permit(
            number="PT-WEB-ACTION-008",
            status=PermitStatus.APPROVED_BY_CHIEF,
        )

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Закрыть")
        self.assertNotContains(response, "Отправить на проверку")


class PermitDocxGenerationViewTests(TestCase):
    """Checks DOCX generation and download endpoints from the permit card."""

    @classmethod
    def setUpTestData(cls):
        call_command("setup_roles", verbosity=0)
        user_model = get_user_model()
        cls.operator = user_model.objects.create_user(username="doc-operator", password="pass")
        cls.manager = user_model.objects.create_user(username="doc-manager", password="pass")
        cls.supervisor = user_model.objects.create_user(username="doc-supervisor", password="pass")
        cls.other_user = user_model.objects.create_user(username="doc-other", password="pass")
        cls.staff_user = user_model.objects.create_user(
            username="doc-staff",
            password="pass",
            is_staff=True,
        )
        cls.operator.groups.add(Group.objects.get(name=ROLE_OPERATOR))
        cls.manager.groups.add(Group.objects.get(name=ROLE_MASTER))
        cls.work_area = WorkArea.objects.create(name="Compressor room")
        cls.equipment = Equipment.objects.create(
            name="Compressor A",
            code="C-100",
            work_area=cls.work_area,
        )
        cls.work_type = WorkType.objects.create(name="Inspection")
        cls.hazard = Hazard.objects.create(name="Stored energy")
        cls.safety_measure = SafetyMeasure.objects.create(name="Isolation")

    def setUp(self):
        from tempfile import TemporaryDirectory

        self.temp_dir = TemporaryDirectory()
        self.override = self.settings(MEDIA_ROOT=self.temp_dir.name)
        self.override.enable()
        self.addCleanup(self.override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.client.force_login(self.operator)

    def make_permit(self, number="PT-WEB-DOCX-001", status=PermitStatus.APPROVED_BY_CHIEF):
        starts_at = timezone.now() + timedelta(days=1)
        return Permit.objects.create(
            number=number,
            status=status,
            work_starts_at=starts_at,
            work_ends_at=starts_at + timedelta(hours=8),
            work_location="Compressor room",
            work_area=self.work_area,
            equipment=self.equipment,
            work_type=self.work_type,
            work_description="Inspect compressor",
            responsible_manager=self.manager,
            work_supervisor=self.supervisor,
            created_by=self.operator,
        )

    def create_template(self):
        from io import BytesIO

        from django.core.files.uploadedfile import SimpleUploadedFile
        from docx import Document

        document = Document()
        document.add_paragraph("Permit {{ permit.number }}")
        document.add_paragraph("Location {{ permit.work_location }}")
        output = BytesIO()
        document.save(output)
        return DocumentTemplate.objects.create(
            name="Permit template",
            document_type="permit",
            version="web-docx-test-1",
            file=SimpleUploadedFile(
                "permit_template.docx",
                output.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            uploaded_by=self.operator,
        )

    def create_generated_document_with_docx(self, number="PT-WEB-PDF-DOCX"):
        permit = self.make_permit(number=number)
        template = self.create_template()
        generated_document = GeneratedDocument.objects.create(
            permit=permit,
            template=template,
            generated_by=self.operator,
        )
        generated_document.file_docx.save(
            "permit.docx",
            ContentFile(b"demo docx"),
            save=True,
        )
        return permit, generated_document

    def assert_download_response_closed(self, response, filename=None):
        try:
            self.assertEqual(response.status_code, 200)
            self.assertIn("attachment", response["Content-Disposition"])
            if filename is not None:
                self.assertIn(filename, response["Content-Disposition"])
            if getattr(response, "streaming", False):
                b"".join(response.streaming_content)
        finally:
            response.close()

    def test_generate_docx_creates_document_and_saves_file(self):
        permit = self.make_permit()
        self.create_template()

        response = self.client.post(
            reverse("permits:generate_docx", kwargs={"pk": permit.pk}),
            follow=True,
        )

        generated_document = GeneratedDocument.objects.get(permit=permit)
        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertTrue(generated_document.file_docx.name.endswith(".docx"))
        self.assertTrue(generated_document.file_docx.storage.exists(generated_document.file_docx.name))
        self.assertEqual(generated_document.file_pdf.name, "")
        self.assertContains(response, "DOCX-документ успешно сформирован")
        self.assertContains(
            response,
            reverse("permits:download_document", kwargs={"pk": generated_document.pk}),
        )

    def test_generated_docx_link_is_shown_on_permit_detail(self):
        permit = self.make_permit(number="PT-WEB-DOCX-002")
        template = self.create_template()
        generated_document = GeneratedDocument.objects.create(
            permit=permit,
            template=template,
            file_docx="generated_documents/docx/permit.docx",
            generated_by=self.operator,
        )

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Сформировать DOCX")
        self.assertContains(
            response,
            reverse("permits:download_document", kwargs={"pk": generated_document.pk}),
        )

    def test_pdf_generation_button_is_shown_when_docx_exists_without_pdf(self):
        permit, generated_document = self.create_generated_document_with_docx("PT-WEB-PDF-BUTTON")

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Сформировать PDF")
        self.assertContains(
            response,
            reverse("permits:generate_pdf", kwargs={"pk": generated_document.pk}),
        )

    @override_settings(PDF_CONVERTER_ENABLED=False)
    def test_generate_pdf_disabled_converter_shows_clear_error(self):
        permit, generated_document = self.create_generated_document_with_docx("PT-WEB-PDF-DISABLED")

        response = self.client.post(
            reverse("permits:generate_pdf", kwargs={"pk": generated_document.pk}),
            follow=True,
        )

        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertContains(response, "Конвертация PDF отключена")
        generated_document.refresh_from_db()
        self.assertEqual(generated_document.file_pdf.name, "")

    @override_settings(PDF_CONVERTER_ENABLED=True, SOFFICE_PATH="missing-soffice")
    @patch("documents.services.shutil.which", return_value=None)
    def test_generate_pdf_missing_soffice_shows_clear_error(self, _which):
        permit, generated_document = self.create_generated_document_with_docx("PT-WEB-PDF-NOSOFFICE")

        response = self.client.post(
            reverse("permits:generate_pdf", kwargs={"pk": generated_document.pk}),
            follow=True,
        )

        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertContains(response, "Исполняемый файл LibreOffice/soffice не найден")
        generated_document.refresh_from_db()
        self.assertEqual(generated_document.file_pdf.name, "")

    @override_settings(PDF_CONVERTER_ENABLED=True, SOFFICE_PATH="soffice")
    @patch("documents.services.shutil.which", return_value="/usr/bin/soffice")
    @patch("documents.services.subprocess.run")
    def test_generate_pdf_successful_mocked_conversion_creates_file_pdf(self, run_mock, _which):
        permit, generated_document = self.create_generated_document_with_docx("PT-WEB-PDF-SUCCESS")

        def fake_run(command, **_kwargs):
            output_dir = Path(command[command.index("--outdir") + 1])
            docx_path = Path(command[-1])
            output_dir.joinpath(f"{docx_path.stem}.pdf").write_bytes(b"%PDF mocked")

        run_mock.side_effect = fake_run

        response = self.client.post(
            reverse("permits:generate_pdf", kwargs={"pk": generated_document.pk}),
            follow=True,
        )

        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertContains(response, "PDF-документ успешно сформирован")
        generated_document.refresh_from_db()
        self.assertTrue(generated_document.file_pdf.name.endswith(".pdf"))
        self.assertContains(
            response,
            reverse("permits:download_pdf", kwargs={"pk": generated_document.pk}),
        )

    def test_generated_pdf_file_can_be_downloaded(self):
        _permit, generated_document = self.create_generated_document_with_docx("PT-WEB-PDF-DOWNLOAD")
        generated_document.file_pdf.save(
            "permit.pdf",
            ContentFile(b"%PDF download"),
            save=True,
        )

        response = self.client.get(reverse("permits:download_pdf", kwargs={"pk": generated_document.pk}))

        try:
            self.assertEqual(response.status_code, 200)
            self.assertIn("attachment", response["Content-Disposition"])
            self.assertIn("permit.pdf", response["Content-Disposition"])
            self.assertEqual(b"".join(response.streaming_content), b"%PDF download")
        finally:
            response.close()

    def test_unauthenticated_user_cannot_download_pdf(self):
        _permit, generated_document = self.create_generated_document_with_docx("PT-WEB-PDF-UNAUTH")
        generated_document.file_pdf.save(
            "permit.pdf",
            ContentFile(b"%PDF download"),
            save=True,
        )
        self.client.logout()

        response = self.client.get(reverse("permits:download_pdf", kwargs={"pk": generated_document.pk}))

        self.assertEqual(response.status_code, 302)
        self.assertIn("/accounts/login/", response["Location"])

    def test_generated_docx_file_can_be_downloaded(self):
        permit = self.make_permit(number="PT-WEB-DOCX-DOWNLOAD")
        self.create_template()
        self.client.post(reverse("permits:generate_docx", kwargs={"pk": permit.pk}))
        generated_document = GeneratedDocument.objects.get(permit=permit)

        response = self.client.get(
            reverse("permits:download_document", kwargs={"pk": generated_document.pk})
        )

        self.assert_download_response_closed(
            response,
            generated_document.file_docx.name.rsplit("/", 1)[-1],
        )

    def test_unauthenticated_user_cannot_download_document(self):
        permit = self.make_permit(number="PT-WEB-DOCX-003")
        template = self.create_template()
        generated_document = GeneratedDocument.objects.create(
            permit=permit,
            template=template,
            file_docx="generated_documents/docx/permit.docx",
            generated_by=self.operator,
        )
        self.client.logout()

        response = self.client.get(reverse("permits:download_document", kwargs={"pk": generated_document.pk}))

        self.assertEqual(response.status_code, 302)
        self.assertIn("/accounts/login/", response["Location"])

    def test_user_with_permit_view_permission_can_download_document(self):
        permit = self.make_permit(number="PT-WEB-DOCX-VIEW-PERM")
        self.create_template()
        self.client.post(reverse("permits:generate_docx", kwargs={"pk": permit.pk}))
        generated_document = GeneratedDocument.objects.get(permit=permit)
        self.client.force_login(self.manager)

        response = self.client.get(
            reverse("permits:download_document", kwargs={"pk": generated_document.pk})
        )

        self.assert_download_response_closed(response)

    def test_unrelated_user_cannot_download_document(self):
        permit = self.make_permit(number="PT-WEB-DOCX-004")
        template = self.create_template()
        generated_document = GeneratedDocument.objects.create(
            permit=permit,
            template=template,
            file_docx="generated_documents/docx/permit.docx",
            generated_by=self.operator,
        )
        self.client.force_login(self.other_user)

        response = self.client.get(reverse("permits:download_document", kwargs={"pk": generated_document.pk}))

        self.assertEqual(response.status_code, 404)

    def test_missing_active_template_shows_clear_error(self):
        permit = self.make_permit(number="PT-WEB-DOCX-005")

        response = self.client.post(
            reverse("permits:generate_docx", kwargs={"pk": permit.pk}),
            follow=True,
        )

        self.assertRedirects(response, reverse("permits:detail", kwargs={"pk": permit.pk}))
        self.assertContains(response, "Активный DOCX-шаблон для нарядов не настроен.")
        self.assertFalse(GeneratedDocument.objects.filter(permit=permit).exists())

    def test_user_without_generation_role_cannot_generate_docx(self):
        permit = self.make_permit(number="PT-WEB-DOCX-NOROLE")
        self.create_template()
        self.client.force_login(self.other_user)

        response = self.client.post(reverse("permits:generate_docx", kwargs={"pk": permit.pk}))

        self.assertEqual(response.status_code, 403)
        self.assertFalse(GeneratedDocument.objects.filter(permit=permit).exists())

    def test_draft_permit_cannot_generate_docx_for_regular_user(self):
        permit = self.make_permit(number="PT-WEB-DOCX-006", status=PermitStatus.DRAFT)
        self.create_template()

        response = self.client.post(reverse("permits:generate_docx", kwargs={"pk": permit.pk}))

        self.assertEqual(response.status_code, 403)
        self.assertFalse(GeneratedDocument.objects.filter(permit=permit).exists())

    def test_staff_user_cannot_generate_docx_for_draft_permit(self):
        permit = self.make_permit(number="PT-WEB-DOCX-007", status=PermitStatus.DRAFT)
        self.create_template()
        self.client.force_login(self.staff_user)

        response = self.client.post(reverse("permits:generate_docx", kwargs={"pk": permit.pk}))

        self.assertEqual(response.status_code, 403)
        self.assertFalse(GeneratedDocument.objects.filter(permit=permit).exists())

    def test_draft_permit_does_not_show_generate_docx_button(self):
        permit = self.make_permit(number="PT-WEB-DOCX-NOBUTTON", status=PermitStatus.DRAFT)

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertNotContains(response, "Сформировать DOCX")
        self.assertContains(response, "Формирование DOCX доступно только для утверждённых или закрытых нарядов")

    def test_closed_permit_shows_generate_docx_button_for_operator(self):
        permit = self.make_permit(number="PT-WEB-DOCX-CLOSED", status=PermitStatus.CLOSED)

        response = self.client.get(reverse("permits:detail", kwargs={"pk": permit.pk}))

        self.assertContains(response, "Сформировать DOCX")
