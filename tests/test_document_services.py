"""Tests for DOCX document generation services."""

from datetime import timedelta
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.utils import timezone
from docx import Document

from documents.models import DocumentTemplate, GeneratedDocument
from documents.services import (
    PdfConversionError,
    _build_permit_context,
    convert_docx_to_pdf,
    generate_permit_docx,
)
from permits.models import (
    Equipment,
    Hazard,
    Permit,
    PermitParticipant,
    PermitParticipantRole,
    Personnel,
    PersonnelGroup,
    SafetyMeasure,
    WorkArea,
    WorkType,
)


class GeneratePermitDocxTests(TestCase):
    """Checks that permit data is rendered into DOCX templates."""

    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.override = override_settings(MEDIA_ROOT=self.temp_dir.name)
        self.override.enable()
        self.addCleanup(self.override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        user_model = get_user_model()
        self.creator = user_model.objects.create_user(username="creator")
        self.manager = user_model.objects.create_user(username="manager")
        self.supervisor = user_model.objects.create_user(username="supervisor")
        self.generator = user_model.objects.create_user(username="generator")
        self.work_area = WorkArea.objects.create(name="Boiler house", description="Main boiler area")
        self.equipment = Equipment.objects.create(
            name="Valve A",
            code="V-100",
            work_area=self.work_area,
            description="Steam valve",
        )
        self.work_type = WorkType.objects.create(name="Repair", description="Repair work")
        self.hazard = Hazard.objects.create(name="Steam", description="Hot steam")
        self.safety_measure = SafetyMeasure.objects.create(name="PPE", description="Wear PPE")
        self.personnel_group = PersonnelGroup.objects.create(name="Masters")
        self.personnel_manager = Personnel.objects.create(
            full_name="John Manager",
            personnel_number="M-001",
            position="Permit master",
            group=self.personnel_group,
            work_area=self.work_area,
        )
        self.personnel_performer = Personnel.objects.create(
            full_name="Paul Performer",
            personnel_number="P-001",
            position="Mechanic",
            group=self.personnel_group,
            work_area=self.work_area,
        )

        starts_at = timezone.now() + timedelta(days=1)
        self.permit = Permit.objects.create(
            number="PT-DOCX-001",
            work_starts_at=starts_at,
            work_ends_at=starts_at + timedelta(hours=8),
            work_location="Boiler room",
            responsible_manager_text="Manual manager",
            work_producer_text="Manual producer",
            work_nature_text="Manual repair nature",
            additional_conditions="Manual extra conditions",
            additional_safety_notes="Manual extra safety notes",
            work_area=self.work_area,
            equipment=self.equipment,
            work_type=self.work_type,
            work_description="Inspect and repair valve",
            responsible_manager=self.manager,
            work_supervisor=self.supervisor,
            created_by=self.creator,
        )
        self.permit.hazards.add(self.hazard)
        self.permit.safety_measures.add(self.safety_measure)
        PermitParticipant.objects.create(
            permit=self.permit,
            role=PermitParticipantRole.RESPONSIBLE_MANAGER,
            personnel=self.personnel_manager,
            note="lead",
            sort_order=1,
        )
        PermitParticipant.objects.create(
            permit=self.permit,
            role=PermitParticipantRole.PERFORMER,
            personnel=self.personnel_performer,
            sort_order=2,
        )
        PermitParticipant.objects.create(
            permit=self.permit,
            role=PermitParticipantRole.BRIGADE_MEMBER,
            manual_name="Manual Brigade Member",
            note="manual",
            sort_order=3,
        )
        self.template = DocumentTemplate.objects.create(
            name="Permit DOCX template",
            document_type="permit",
            version="docx-test-1",
            file=SimpleUploadedFile(
                "permit_template.docx",
                self._template_bytes(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            uploaded_by=self.creator,
        )

    def test_generate_permit_docx_renders_template_and_saves_generated_document(self):
        generated_document = generate_permit_docx(
            permit_id=self.permit.pk,
            template_id=self.template.pk,
            user=self.generator,
        )

        self.assertIsInstance(generated_document, GeneratedDocument)
        self.assertEqual(generated_document.permit, self.permit)
        self.assertEqual(generated_document.template, self.template)
        self.assertEqual(generated_document.generated_by, self.generator)
        self.assertTrue(generated_document.file_docx.name.endswith(".docx"))
        self.assertEqual(generated_document.file_pdf.name, "")

        rendered_doc = Document(generated_document.file_docx.path)
        rendered_text = "\n".join(paragraph.text for paragraph in rendered_doc.paragraphs)
        self.assertIn("Permit PT-DOCX-001", rendered_text)
        self.assertIn("Location Boiler room", rendered_text)
        self.assertIn("Supervisor supervisor", rendered_text)
        self.assertIn("Description Inspect and repair valve", rendered_text)
        self.assertIn("Work area Boiler house", rendered_text)
        self.assertIn("Equipment Valve A V-100", rendered_text)
        self.assertIn("Work type Repair", rendered_text)
        self.assertIn("Hazards Steam", rendered_text)
        self.assertIn("Safety PPE", rendered_text)
        self.assertIn("Nomer PT-DOCX-001", rendered_text)
        self.assertIn("Status Черновик", rendered_text)
        self.assertIn("Uchastok Boiler house", rendered_text)
        self.assertIn("Oborudovanie Valve A", rendered_text)
        self.assertIn("Vid Repair", rendered_text)
        self.assertIn("Mesto Boiler room", rendered_text)
        self.assertIn("Opisanie Inspect and repair valve", rendered_text)
        self.assertIn(f"Data start {timezone.localtime(self.permit.work_starts_at):%d.%m.%Y}", rendered_text)
        self.assertIn(f"Data end {timezone.localtime(self.permit.work_ends_at):%d.%m.%Y}", rendered_text)
        self.assertIn(f"Data created {timezone.localtime(self.permit.created_at):%d.%m.%Y}", rendered_text)
        self.assertIn("Rukovoditel John Manager", rendered_text)
        self.assertIn("Ispolniteli Paul Performer", rendered_text)
        self.assertIn("Brigada Manual Brigade Member", rendered_text)
        self.assertIn("Uchastniki John Manager", rendered_text)
        self.assertIn("Proizvoditel Manual producer", rendered_text)
        self.assertIn("Harakter Manual repair nature", rendered_text)
        self.assertIn("Usloviya Manual extra conditions", rendered_text)
        self.assertIn("Dop mery Manual extra safety notes", rendered_text)
        self.assertIn("Sozdal creator", rendered_text)
        self.assertIn("Opasnosti Steam", rendered_text)
        self.assertIn("Mery PPE", rendered_text)
        self.assertIn("Shablon Permit DOCX template", rendered_text)
        self.assertIn("Versiya docx-test-1", rendered_text)

    def test_build_permit_context_contains_reference_directory_fields(self):
        context = _build_permit_context(self.permit)["permit"]

        self.assertEqual(context["work_area_name"], "Boiler house")
        self.assertEqual(context["equipment_name"], "Valve A")
        self.assertEqual(context["equipment_code"], "V-100")
        self.assertEqual(context["work_type_name"], "Repair")
        self.assertEqual(context["hazard_names"], ["Steam"])
        self.assertEqual(context["hazard_names_text"], "Steam")
        self.assertEqual(context["safety_measure_names"], ["PPE"])
        self.assertEqual(context["safety_measure_names_text"], "PPE")

    def test_build_permit_context_contains_translit_docx_aliases(self):
        context = _build_permit_context(self.permit, self.template)

        self.assertEqual(context["nomer_naryada"], "PT-DOCX-001")
        self.assertEqual(context["status_naryada"], "Черновик")
        self.assertEqual(context["uchastok"], "Boiler house")
        self.assertEqual(context["oborudovanie"], "Valve A")
        self.assertEqual(context["vid_rabot"], "Repair")
        self.assertEqual(context["mesto_rabot"], "Boiler room")
        self.assertEqual(context["opisanie_rabot"], "Inspect and repair valve")
        self.assertEqual(context["data_nachala"], timezone.localtime(self.permit.work_starts_at).strftime("%d.%m.%Y"))
        self.assertEqual(context["data_okonchaniya"], timezone.localtime(self.permit.work_ends_at).strftime("%d.%m.%Y"))
        self.assertEqual(context["data_sozdaniya"], timezone.localtime(self.permit.created_at).strftime("%d.%m.%Y"))
        self.assertEqual(context["otvetstvennyy_rukovoditel"], "John Manager — Permit master — Masters (lead)")
        self.assertEqual(context["proizvoditel_rabot"], "Manual producer")
        self.assertEqual(context["harakter_rabot"], "Manual repair nature")
        self.assertEqual(context["dopolnitelnye_usloviya"], "Manual extra conditions")
        self.assertEqual(context["dopolnitelnye_mery_bezopasnosti"], "Manual extra safety notes")
        self.assertEqual(context["sozdal_polzovatel"], "creator")
        self.assertEqual(context["opasnosti"], "Steam")
        self.assertEqual(context["mery_bezopasnosti"], "PPE")
        self.assertEqual(context["otvetstvennye_rukovoditeli"], "John Manager — Permit master — Masters (lead)")
        self.assertEqual(context["ispolniteli"], "Paul Performer — Mechanic — Masters")
        self.assertEqual(context["chleny_brigady"], "Manual Brigade Member (manual)")
        self.assertIn("John Manager — Permit master — Masters (lead)", context["uchastniki_rabot"])
        self.assertIn("Manual Brigade Member (manual)", context["uchastniki_rabot"])
        self.assertEqual(context["shablon_dokumenta"], "Permit DOCX template")
        self.assertEqual(context["versiya_shablona"], "docx-test-1")

    @override_settings(PDF_CONVERTER_ENABLED=False)
    def test_convert_docx_to_pdf_reports_disabled_converter(self):
        generated_document = self._generated_document_with_docx()

        with self.assertRaisesMessage(PdfConversionError, "Конвертация PDF отключена"):
            convert_docx_to_pdf(generated_document.pk)

    @override_settings(PDF_CONVERTER_ENABLED=True, SOFFICE_PATH="missing-soffice")
    @patch("documents.services.shutil.which", return_value=None)
    def test_convert_docx_to_pdf_reports_missing_soffice(self, _which):
        generated_document = self._generated_document_with_docx()

        with self.assertRaisesMessage(PdfConversionError, "Исполняемый файл LibreOffice/soffice не найден"):
            convert_docx_to_pdf(generated_document.pk)

    @override_settings(PDF_CONVERTER_ENABLED=True, SOFFICE_PATH="soffice")
    @patch("documents.services.shutil.which", return_value="/usr/bin/soffice")
    @patch("documents.services.subprocess.run")
    def test_convert_docx_to_pdf_saves_pdf_with_mocked_soffice(self, run_mock, _which):
        generated_document = self._generated_document_with_docx()

        def fake_run(command, **_kwargs):
            output_dir = Path(command[command.index("--outdir") + 1])
            docx_path = Path(command[-1])
            output_dir.joinpath(f"{docx_path.stem}.pdf").write_bytes(b"%PDF-1.4 demo")

        run_mock.side_effect = fake_run

        converted_document = convert_docx_to_pdf(generated_document.pk)

        self.assertEqual(converted_document.pk, generated_document.pk)
        self.assertTrue(converted_document.file_pdf.name.endswith(".pdf"))
        with converted_document.file_pdf.open("rb") as pdf_file:
            self.assertEqual(pdf_file.read(), b"%PDF-1.4 demo")
        run_mock.assert_called_once()

    def _template_bytes(self):
        document = Document()
        document.add_paragraph("Permit {{ permit.number }}")
        document.add_paragraph("Location {{ permit.work_location }}")
        document.add_paragraph("Supervisor {{ permit.work_supervisor_name }}")
        document.add_paragraph("Description {{ permit.work_description }}")
        document.add_paragraph("Work area {{ permit.work_area_name }}")
        document.add_paragraph("Equipment {{ permit.equipment_name }} {{ permit.equipment_code }}")
        document.add_paragraph("Work type {{ permit.work_type_name }}")
        document.add_paragraph("Hazards {{ permit.hazard_names_text }}")
        document.add_paragraph("Safety {{ permit.safety_measure_names_text }}")
        document.add_paragraph("Nomer {{ nomer_naryada }}")
        document.add_paragraph("Status {{ status_naryada }}")
        document.add_paragraph("Uchastok {{ uchastok }}")
        document.add_paragraph("Oborudovanie {{ oborudovanie }}")
        document.add_paragraph("Vid {{ vid_rabot }}")
        document.add_paragraph("Mesto {{ mesto_rabot }}")
        document.add_paragraph("Opisanie {{ opisanie_rabot }}")
        document.add_paragraph("Data start {{ data_nachala }}")
        document.add_paragraph("Data end {{ data_okonchaniya }}")
        document.add_paragraph("Data created {{ data_sozdaniya }}")
        document.add_paragraph("Harakter {{ harakter_rabot }}")
        document.add_paragraph("Usloviya {{ dopolnitelnye_usloviya }}")
        document.add_paragraph("Dop mery {{ dopolnitelnye_mery_bezopasnosti }}")
        document.add_paragraph("Rukovoditel {{ otvetstvennyy_rukovoditel }}")
        document.add_paragraph("Proizvoditel {{ proizvoditel_rabot }}")
        document.add_paragraph("Ispolniteli {{ ispolniteli }}")
        document.add_paragraph("Brigada {{ chleny_brigady }}")
        document.add_paragraph("Uchastniki {{ uchastniki_rabot }}")
        document.add_paragraph("Sozdal {{ sozdal_polzovatel }}")
        document.add_paragraph("Opasnosti {{ opasnosti }}")
        document.add_paragraph("Mery {{ mery_bezopasnosti }}")
        document.add_paragraph("Shablon {{ shablon_dokumenta }}")
        document.add_paragraph("Versiya {{ versiya_shablona }}")
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _generated_document_with_docx(self):
        generated_document = GeneratedDocument.objects.create(
            permit=self.permit,
            template=self.template,
            generated_by=self.generator,
        )
        generated_document.file_docx.save(
            "generated-test.docx",
            ContentFile(b"demo docx"),
            save=True,
        )
        return generated_document
