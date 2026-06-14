"""Tests for DocumentTemplate admin helper pages."""

from io import BytesIO
from tempfile import TemporaryDirectory

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from documents.forms import DOCX_TEMPLATE_ERROR_MESSAGE
from documents.models import DocumentTemplate, GeneratedDocument
from permits.models import Permit


class DocumentTemplateAdminHelpTests(TestCase):
    """Checks that DOCX template variable help is available in admin."""

    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.override = override_settings(MEDIA_ROOT=self.temp_dir.name)
        self.override.enable()
        self.addCleanup(self.override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.admin_user = get_user_model().objects.create_superuser(
            username="admin",
            email="admin@example.com",
            password="password",
        )
        self.client.force_login(self.admin_user)

    def test_template_variables_help_page_is_available(self):
        response = self.client.get(reverse("admin:documents_documenttemplate_template_variables"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Переменные DOCX-шаблона")
        self.assertContains(response, "{{ nomer_naryada }}")
        self.assertContains(response, "{{ mery_bezopasnosti }}")
        self.assertContains(response, "{{ dopolnitelnye_usloviya }}")
        self.assertContains(response, "использовать только транслит-переменные")

    def test_document_template_changelist_links_to_variable_help_and_test_button(self):
        template = self._template("valid.docx", "Наряд {{ nomer_naryada }}")

        response = self.client.get(reverse("admin:documents_documenttemplate_changelist"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Переменные шаблона")
        self.assertContains(response, "Проверить шаблон")
        self.assertContains(response, reverse("admin:documents_documenttemplate_template_variables"))
        self.assertContains(response, reverse("admin:documents_documenttemplate_test", args=[template.pk]))

    def test_valid_template_can_be_checked_and_downloaded_without_generated_document(self):
        template = self._template(
            "valid.docx",
            "Наряд {{ nomer_naryada }} {{ mesto_rabot }} {{ dopolnitelnye_usloviya }}",
        )

        response = self.client.get(reverse("admin:documents_documenttemplate_test", args=[template.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment", response["Content-Disposition"])
        self.assertIn(".docx", response["Content-Disposition"])
        self.assertEqual(GeneratedDocument.objects.count(), 0)
        self.assertEqual(Permit.objects.count(), 0)

        rendered_doc = Document(BytesIO(response.content))
        rendered_text = "\n".join(paragraph.text for paragraph in rendered_doc.paragraphs)
        self.assertIn("Наряд DEMO-001", rendered_text)
        self.assertIn("Площадка обслуживания насосной станции", rendered_text)
        self.assertIn("Работы выполнять после инструктажа", rendered_text)

    def test_broken_template_check_shows_clear_error_without_generated_document(self):
        template = self._template("broken.docx", "Наряд {{ } }")

        response = self.client.get(reverse("admin:documents_documenttemplate_test", args=[template.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, DOCX_TEMPLATE_ERROR_MESSAGE)
        self.assertContains(response, "Ошибка проверки DOCX-шаблона")
        self.assertEqual(GeneratedDocument.objects.count(), 0)
        self.assertEqual(Permit.objects.count(), 0)

    def _template(self, filename, paragraph_text):
        return DocumentTemplate.objects.create(
            name=f"Template {filename}",
            document_type="permit",
            version=filename,
            file=SimpleUploadedFile(
                filename,
                self._docx_bytes(paragraph_text),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            uploaded_by=self.admin_user,
        )

    def _docx_bytes(self, paragraph_text):
        document = Document()
        document.add_paragraph(paragraph_text)
        output = BytesIO()
        document.save(output)
        return output.getvalue()
